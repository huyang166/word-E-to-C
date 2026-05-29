from pathlib import Path

from docx import Document
from docx.text.paragraph import Paragraph

from app.models import BlockKind, Side, TextBlock


class DocxService:
    def extract_blocks(self, path: Path, side: Side) -> list[TextBlock]:
        document = Document(path)
        blocks: list[TextBlock] = []

        for paragraph_index, paragraph in enumerate(document.paragraphs):
            text = paragraph.text.strip()
            if not text:
                continue
            blocks.append(
                self._block_from_paragraph(
                    paragraph=paragraph,
                    side=side,
                    index=len(blocks),
                    path=f"p:{paragraph_index}",
                    kind=self._paragraph_kind(paragraph),
                )
            )

        for table_index, table in enumerate(document.tables):
            for row_index, row in enumerate(table.rows):
                for cell_index, cell in enumerate(row.cells):
                    for paragraph_index, paragraph in enumerate(cell.paragraphs):
                        text = paragraph.text.strip()
                        if not text:
                            continue
                        blocks.append(
                            self._block_from_paragraph(
                                paragraph=paragraph,
                                side=side,
                                index=len(blocks),
                                path=f"table:{table_index}:row:{row_index}:cell:{cell_index}:p:{paragraph_index}",
                                kind=BlockKind.TABLE_CELL,
                            )
                        )

        return blocks

    def export_with_replacements(
        self,
        source_path: Path,
        output_path: Path,
        replacements: dict[str, str],
    ) -> None:
        document = Document(source_path)
        for block_path, replacement_text in replacements.items():
            paragraph = self._resolve_paragraph(document, block_path)
            self._replace_paragraph_text(paragraph, replacement_text)
        output_path.parent.mkdir(parents=True, exist_ok=True)
        document.save(output_path)

    def _resolve_paragraph(self, document, block_path: str) -> Paragraph:
        parts = block_path.split(":")
        if parts[0] == "p":
            return document.paragraphs[int(parts[1])]
        if parts[0] == "table":
            table_index = int(parts[1])
            row_index = int(parts[3])
            cell_index = int(parts[5])
            paragraph_index = int(parts[7])
            return document.tables[table_index].rows[row_index].cells[cell_index].paragraphs[paragraph_index]
        raise ValueError(f"Unsupported block path: {block_path}")

    def _replace_paragraph_text(self, paragraph: Paragraph, text: str) -> None:
        if paragraph.runs:
            paragraph.runs[0].text = text
            for run in paragraph.runs[1:]:
                run.text = ""
            return
        paragraph.add_run(text)

    def _block_from_paragraph(
        self,
        paragraph: Paragraph,
        side: Side,
        index: int,
        path: str,
        kind: BlockKind,
    ) -> TextBlock:
        return TextBlock(
            id=f"{side.value}-{index:05d}",
            side=side,
            kind=kind,
            index=index,
            text=paragraph.text.strip(),
            path=path,
            mappedId=None,
        )

    def _paragraph_kind(self, paragraph: Paragraph) -> BlockKind:
        style_name = paragraph.style.name.lower() if paragraph.style and paragraph.style.name else ""
        if style_name.startswith("heading"):
            return BlockKind.HEADING
        return BlockKind.PARAGRAPH
