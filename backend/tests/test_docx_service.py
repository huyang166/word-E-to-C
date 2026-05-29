from app.docx_service import DocxService
from app.models import BlockKind, Side
from docx import Document


def test_extracts_headings_body_paragraphs_and_table_cells(sample_docx):
    service = DocxService()

    blocks = service.extract_blocks(sample_docx, Side.EN)

    assert [block.text for block in blocks] == [
        "Introduction",
        "First body paragraph.",
        "Table cell text.",
    ]
    assert [block.kind for block in blocks] == [
        BlockKind.HEADING,
        BlockKind.PARAGRAPH,
        BlockKind.TABLE_CELL,
    ]
    assert blocks[0].id == "en-00000"
    assert blocks[2].path == "table:0:row:0:cell:0:p:0"


def test_replaces_body_paragraph_and_preserves_other_text(sample_docx, tmp_path):
    service = DocxService()
    output = tmp_path / "updated.docx"

    service.export_with_replacements(
        source_path=sample_docx,
        output_path=output,
        replacements={"p:1": "Updated body paragraph."},
    )

    doc = Document(output)
    assert doc.paragraphs[0].text == "Introduction"
    assert doc.paragraphs[1].text == "Updated body paragraph."
    assert doc.tables[0].cell(0, 0).text == "Table cell text."


def test_replaces_table_cell_paragraph(sample_docx, tmp_path):
    service = DocxService()
    output = tmp_path / "updated-table.docx"

    service.export_with_replacements(
        source_path=sample_docx,
        output_path=output,
        replacements={"table:0:row:0:cell:0:p:0": "Updated table text."},
    )

    doc = Document(output)
    assert doc.tables[0].cell(0, 0).text == "Updated table text."
