from pathlib import Path

from docx import Document
from fastapi.testclient import TestClient

from app.config import Settings
from app.main import create_app


def make_docx(path: Path, heading: str, body: str) -> None:
    doc = Document()
    doc.add_heading(heading, level=1)
    doc.add_paragraph(body)
    doc.save(path)


def test_health_endpoint_returns_ok(tmp_path):
    app = create_app(Settings(data_dir=tmp_path, openai_api_key="", openai_model=""))
    client = TestClient(app)

    response = client.get("/api/health")

    assert response.status_code == 200
    assert response.json() == {"status": "ok"}


def test_uploads_and_parses_paired_docx(tmp_path):
    en_path = tmp_path / "en.docx"
    zh_path = tmp_path / "zh.docx"
    make_docx(en_path, "Introduction", "English body.")
    make_docx(zh_path, "引言", "中文正文。")
    app = create_app(Settings(data_dir=tmp_path / "projects", openai_api_key="", openai_model=""))
    client = TestClient(app)

    with en_path.open("rb") as en_file, zh_path.open("rb") as zh_file:
        response = client.post(
            "/api/projects",
            files={
                "en_file": (
                    "en.docx",
                    en_file,
                    "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                ),
                "zh_file": (
                    "zh.docx",
                    zh_file,
                    "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                ),
            },
        )

    assert response.status_code == 200
    payload = response.json()
    assert payload["projectId"]
    assert payload["enBlocks"][0]["text"] == "Introduction"
    assert payload["zhBlocks"][1]["text"] == "中文正文。"
    assert payload["mappings"][0]["enId"] == "en-00000"
    assert payload["mappings"][0]["zhId"] == "zh-00000"


def test_rejects_non_docx_uploads(tmp_path):
    txt_path = tmp_path / "bad.txt"
    txt_path.write_text("not a docx", encoding="utf-8")
    docx_path = tmp_path / "zh.docx"
    make_docx(docx_path, "引言", "正文")
    app = create_app(Settings(data_dir=tmp_path / "projects", openai_api_key="", openai_model=""))
    client = TestClient(app)

    with txt_path.open("rb") as en_file, docx_path.open("rb") as zh_file:
        response = client.post(
            "/api/projects",
            files={
                "en_file": ("bad.txt", en_file, "text/plain"),
                "zh_file": (
                    "zh.docx",
                    zh_file,
                    "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                ),
            },
        )

    assert response.status_code == 400
    assert response.json()["detail"] == "请上传 .docx 格式的 Word 文件。"


def test_suggest_without_api_key_returns_chinese_error(tmp_path):
    app = create_app(Settings(data_dir=tmp_path / "projects", openai_api_key="", openai_model="test-model"))
    client = TestClient(app)

    response = client.post(
        "/api/projects/demo/suggest",
        json={
            "direction": "en_to_zh",
            "sourceBlockId": "en-00000",
            "targetBlockId": "zh-00000",
            "sourceText": "Updated English.",
            "targetText": "原中文。",
        },
    )

    assert response.status_code == 400
    assert response.json()["detail"] == "未配置 API Key，请在 .env 中设置。"
