from pathlib import Path

import pytest
from docx import Document


@pytest.fixture
def sample_docx(tmp_path: Path) -> Path:
    path = tmp_path / "sample.docx"
    doc = Document()
    doc.add_heading("Introduction", level=1)
    doc.add_paragraph("First body paragraph.")
    table = doc.add_table(rows=1, cols=1)
    table.cell(0, 0).text = "Table cell text."
    doc.save(path)
    return path
